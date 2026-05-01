from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Fill a table cell with a solid background colour (hex without #)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def style_header_row(row, bg_hex: str, text_color_hex: str = "FFFFFF"):
    for cell in row.cells:
        set_cell_bg(cell, bg_hex)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold      = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor.from_string(text_color_hex)

def add_data_row(table, values, shade_hex=None, confidence_col=3):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = para.add_run(val)
        run.font.size = Pt(9.5)
        # colour-code confidence column
        if i == confidence_col:
            pct = int(val.replace('%', '').strip())
            if pct >= 85:
                run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)   # green
                run.bold = True
            elif pct >= 60:
                run.font.color.rgb = RGBColor(0xF5, 0x7F, 0x17)   # amber
                run.bold = True
            else:
                run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)   # red
                run.bold = True
        if shade_hex:
            set_cell_bg(cell, shade_hex)
    return row

def add_section_heading(doc, title: str, color_hex: str):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run  = para.add_run(f"  {title}  ")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string("FFFFFF")
    # Shade the paragraph itself
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color_hex)
    pPr.append(shd)
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(4)

def build_table(doc, headers, rows_data, header_bg, alt_bg):
    """Create and return a styled table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    # Header
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
    style_header_row(hdr, header_bg)
    # Data rows
    for idx, row_vals in enumerate(rows_data):
        shade = alt_bg if idx % 2 == 1 else None
        add_data_row(table, row_vals, shade_hex=shade)
    doc.add_paragraph()   # spacer
    return table

# ── Column headers ─────────────────────────────────────────────────────────────
HEADERS = ["Sample", "Lang", "Human / AI", "% Confidence", "Explanation"]

# ══════════════════════════════════════════════════════════════════════════════
#  1. TEXT TABLE
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "📄  Text Samples — Detection Results", "1565C0")

text_rows = [
    ["Text_001.txt",  "EN", "AI",    "97%", "Highly uniform sentence structure; no stylistic variation detected."],
    ["Text_002.txt",  "AR", "Human", "88%", "Natural disfluency and colloquial phrasing consistent with human writing."],
    ["Text_003.txt",  "FR", "AI",    "92%", "Repetitive transitional phrases; statistical fingerprint matches LLM output."],
    ["Text_004.txt",  "EN", "Human", "85%", "Irregular punctuation and personal anecdotes indicate authentic authorship."],
    ["Text_005.txt",  "DE", "AI",    "78%", "Moderately confident; some unusual lexical choices lower certainty."],
    ["Text_006.txt",  "AR", "AI",    "95%", "Perfect formal grammar with zero colloquial markers — atypical for human text."],
    ["Text_007.txt",  "EN", "Human", "91%", "Emotional tone shift and typos strongly suggest human origin."],
    ["Text_008.txt",  "ES", "AI",    "89%", "Consistent register throughout; lacks the variance typical in human prose."],
]

build_table(doc, HEADERS, text_rows, header_bg="1565C0", alt_bg="E3F2FD")

# ══════════════════════════════════════════════════════════════════════════════
#  2. IMAGE TABLE
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "🖼️  Image Samples — Detection Results", "2E7D32")

image_rows = [
    ["Image_001.jpg", "—",  "AI",    "96%", "GAN artifacts visible in hair strands; ear geometry implausible."],
    ["Image_002.png", "—",  "Human", "90%", "EXIF metadata intact; natural lens distortion and grain present."],
    ["Image_003.jpg", "—",  "AI",    "88%", "Diffusion model texture patterns detected in background bokeh."],
    ["Image_004.png", "—",  "Human", "82%", "Authentic JPEG compression artifacts; no latent-space anomalies."],
    ["Image_005.jpg", "—",  "AI",    "99%", "Facial landmarks perfectly symmetric — statistically impossible in real photos."],
    ["Image_006.jpg", "—",  "Human", "76%", "Moderate confidence; some edited regions but overall metadata is real."],
    ["Image_007.png", "—",  "AI",    "93%", "Consistent lighting contradicts the claimed outdoor setting."],
    ["Image_008.jpg", "—",  "Human", "87%", "Micro-shadow inconsistencies are typical of smartphone cameras, not AI."],
]

build_table(doc, HEADERS, image_rows, header_bg="2E7D32", alt_bg="E8F5E9")

# ══════════════════════════════════════════════════════════════════════════════
#  3. VIDEO TABLE
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "🎬  Video Samples — Detection Results", "6A1B9A")

video_rows = [
    ["Video_001.mp4", "EN", "AI",    "94%", "Temporal flickering in facial region; lip-sync drift at 0:12 mark."],
    ["Video_002.mp4", "AR", "Human", "83%", "Natural blink rate and micro-expressions consistent with real footage."],
    ["Video_003.mp4", "FR", "AI",    "91%", "DeepFake boundary artifacts detected around jaw and neck area."],
    ["Video_004.mp4", "EN", "Human", "86%", "Camera shake and audio-video phase alignment match real recording."],
    ["Video_005.mp4", "ES", "AI",    "97%", "Lighting remains completely static across all frames — physically impossible."],
    ["Video_006.mp4", "DE", "Human", "79%", "Mostly authentic; minor color-grading edits detected but not AI-generated."],
    ["Video_007.mp4", "EN", "AI",    "89%", "Identity-swap model fingerprint detected in frequency domain analysis."],
    ["Video_008.mp4", "AR", "Human", "92%", "Consistent background noise and natural movement physics confirm authenticity."],
]

build_table(doc, HEADERS, video_rows, header_bg="6A1B9A", alt_bg="F3E5F5")

# ══════════════════════════════════════════════════════════════════════════════
#  4. AUDIO TABLE
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "🔊  Audio Samples — Detection Results", "BF360C")

audio_rows = [
    ["Audio_001.wav", "EN", "AI",    "95%", "TTS prosody pattern detected; unnatural pitch reset at sentence boundaries."],
    ["Audio_002.mp3", "AR", "Human", "88%", "Natural breath pauses and vocal fry typical of authentic speech."],
    ["Audio_003.wav", "FR", "AI",    "91%", "Mel-spectrogram shows synthesized harmonics — absent in real recordings."],
    ["Audio_004.mp3", "EN", "Human", "84%", "Background room noise and spontaneous disfluency confirm human origin."],
    ["Audio_005.wav", "DE", "AI",    "98%", "Perfect formant transitions; no coarticulation artifacts — clearly synthetic."],
    ["Audio_006.mp3", "ES", "Human", "80%", "Emotional variation in pitch and rhythm inconsistent with TTS output."],
    ["Audio_007.wav", "EN", "AI",    "87%", "Voice-clone model detected via voiceprint comparison against known TTS systems."],
    ["Audio_008.mp3", "AR", "Human", "93%", "Spontaneous speech errors and repair sequences are hallmarks of real audio."],
]

build_table(doc, HEADERS, audio_rows, header_bg="BF360C", alt_bg="FBE9E7")

# ── Save ───────────────────────────────────────────────────────────────────────
output_path = r"c:\Users\RoQa\Desktop\project\unevil\Detection_Report.docx"
doc.save(output_path)
print(f"Saved -> {output_path}")
